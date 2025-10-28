"""
Document Analyzer - Analisi completa documenti Word
"""

from docx import Document
from typing import Dict, List, Optional
import re


class WordAnalyzer:
    def __init__(self, knowledge_base):
        self.kb = knowledge_base
        self.document = None
        self.doc_path = None
        
        # Inizializza classificatore formule
        try:
            from core.formula_classifier import FormulaClassifier
            self.formula_classifier = FormulaClassifier()
        except:
            self.formula_classifier = None
        
    def load_document(self, doc_path: str):
        """Carica documento Word"""
        self.doc_path = doc_path
        self.document = Document(doc_path)
        
    def analyze(self, discipline: str = "mathematics") -> Dict:
        """Analizza il documento"""
        if not self.document:
            raise ValueError("Nessun documento caricato")
        
        # Crea mappa delle pagine una volta sola per tutti gli elementi
        page_map = self._build_page_map()
        
        results = {
            'images': self._extract_images(page_map),
            'tables': self._extract_tables(page_map),
            'equations': self._extract_equations(page_map),
            'references': self._extract_references(page_map)
        }
        
        return results
    
    def _extract_images(self, page_map: Dict[int, int]) -> List[Dict]:
        """Estrae immagini con contesto 5 righe sopra/sotto"""
        images = []
        
        for rel_id, rel in self.document.part.rels.items():
            if "image" in rel.target_ref:
                location = self._find_image_paragraph(rel_id)
                
                if location is not None:
                    # Ottieni contesto 5 righe
                    context = self._get_context_5_lines(location)
                    
                    # Cerca label
                    label = self._extract_image_label(
                        ' '.join(context['before']) + ' ' + ' '.join(context['after'])
                    )
                    
                    # Rileva elementi adiacenti
                    adjacent = self._detect_adjacent_elements(location)
                    
                    # Ottieni numero di pagina reale
                    page_number = page_map.get(location, 1)
                    
                    image_data = {
                        'paragraph_index': location,
                        'page_number': page_number,  # NUOVO: numero di pagina reale
                        'label': label or f"IMG_{len(images)+1}",
                        'rel_id': rel_id,
                        'filename': rel.target_ref.split('/')[-1],
                        'content_type': rel.target_part.content_type,
                        'image_part': rel.target_part,
                        'context_before': context['before'],  # Lista di 5 stringhe
                        'context_after': context['after'],    # Lista di 5 stringhe
                        'has_adjacent_before': adjacent['before'],
                        'has_adjacent_after': adjacent['after']
                    }
                    
                    images.append(image_data)
        
        return images
    
    def _find_image_paragraph(self, rel_id: str) -> Optional[int]:
        """Trova paragrafo contenente immagine"""
        for i, paragraph in enumerate(self.document.paragraphs):
            for run in paragraph.runs:
                run_element = run._element
                blips = run_element.findall('.//{*}blip')
                for blip in blips:
                    embed_id = blip.get('{http://schemas.openxmlformats.org/officeDocument/2006/relationships}embed')
                    if embed_id == rel_id:
                        return i
        return None
    
    def _get_context_5_lines(self, para_idx: int, context_lines: int = 10) -> Dict:
        """
        Ottiene righe di contesto prima e dopo un paragrafo.
        
        Args:
            para_idx: Indice del paragrafo
            context_lines: Numero di righe di contesto (default 10 per lato)
            
        Returns:
            Dict con 'before' e 'after' contenenti le righe di contesto
        """
        before_lines = []
        after_lines = []
        
        # N paragrafi prima
        for i in range(max(0, para_idx - context_lines), para_idx):
            before_lines.append(self.document.paragraphs[i].text)
        
        # N paragrafi dopo
        for i in range(para_idx + 1, min(len(self.document.paragraphs), para_idx + context_lines + 1)):
            after_lines.append(self.document.paragraphs[i].text)
        
        return {
            'before': before_lines,
            'after': after_lines
        }
    
    def _detect_adjacent_elements(self, para_idx: int) -> Dict:
        """Rileva se ci sono elementi adiacenti (immagini/tabelle/equazioni) sopra/sotto"""
        adjacent = {'before': False, 'after': False}
        
        # Controlla paragrafo precedente
        if para_idx > 0:
            prev_para = self.document.paragraphs[para_idx - 1]
            if self._has_special_elements(prev_para):
                adjacent['before'] = True
        
        # Controlla paragrafo successivo
        if para_idx < len(self.document.paragraphs) - 1:
            next_para = self.document.paragraphs[para_idx + 1]
            if self._has_special_elements(next_para):
                adjacent['after'] = True
        
        return adjacent
    
    def _has_special_elements(self, paragraph) -> bool:
        """Verifica se paragrafo contiene immagini o equazioni"""
        # Controlla immagini
        for run in paragraph.runs:
            blips = run._element.findall('.//{*}blip')
            if blips:
                return True
        
        # Controlla equazioni
        omath = paragraph._element.findall('.//{*}oMath')
        if omath:
            return True
        
        return False
    
    def _build_page_map(self) -> Dict[int, int]:
        """
        Costruisce una mappa dei numeri di pagina per ogni paragrafo.
        
        NOTA: python-docx non fornisce accesso diretto ai numeri di pagina.
        Questo metodo usa una stima avanzata basata su:
        - Page breaks espliciti (quando presenti)
        - Altezza stimata dei paragrafi
        - Presenza di immagini
        - Stili dei paragrafi (heading, normal)
        
        Returns:
            Dict[int, int]: Dizionario {paragraph_index: page_number}
        """
        page_map = {}
        current_page = 1
        current_page_height = 0  # Altezza accumulata nella pagina corrente (in righe)
        
        # Altezze standard (in righe equivalenti)
        MAX_PAGE_HEIGHT = 50  # ~50 righe per pagina A4 standard
        LINE_HEIGHT = 1  # 1 riga normale
        HEADING_HEIGHT = 2  # Heading conta come 2 righe
        EMPTY_LINE = 0.5  # Paragrafo vuoto
        IMAGE_HEIGHT = 15  # Immagine media ~ 15 righe
        
        # Namespace per XML
        w_namespace = '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}'
        
        for i, paragraph in enumerate(self.document.paragraphs):
            # Assegna il paragrafo alla pagina corrente
            page_map[i] = current_page
            
            # === CONTROLLO PAGE BREAKS ESPLICITI ===
            para_element = paragraph._element
            
            # Cerca page breaks nei runs
            for run in paragraph.runs:
                run_element = run._element
                page_breaks = run_element.findall(f'.//{w_namespace}br[@{w_namespace}type="page"]')
                if page_breaks:
                    current_page += len(page_breaks)
                    current_page_height = 0  # Reset altezza dopo page break
                    page_map[i] = current_page  # Aggiorna la pagina per questo paragrafo
            
            # Cerca section breaks
            section_props = para_element.findall(f'.//{w_namespace}sectPr')
            if section_props:
                for sect in section_props:
                    page_break_type = sect.find(f'{w_namespace}type')
                    if page_break_type is not None:
                        type_val = page_break_type.get(f'{w_namespace}val')
                        if type_val in ['nextPage', 'evenPage', 'oddPage']:
                            current_page += 1
                            current_page_height = 0
                            page_map[i] = current_page
            
            # === STIMA ALTEZZA PARAGRAFO ===
            para_height = 0
            
            # Paragrafo vuoto
            if not paragraph.text.strip():
                para_height = EMPTY_LINE
            
            # Heading (conta di più)
            elif paragraph.style.name.startswith('Heading'):
                para_height = HEADING_HEIGHT
            
            # Paragrafo con immagine
            elif self._paragraph_has_image(paragraph):
                para_height = IMAGE_HEIGHT
            
            # Paragrafo normale - stima in base al numero di righe
            else:
                # Stima righe in base alla lunghezza del testo
                # ~90 caratteri per riga in media
                chars = len(paragraph.text)
                estimated_lines = max(1, chars / 90)
                para_height = estimated_lines
            
            # Aggiungi altezza alla pagina corrente
            current_page_height += para_height
            
            # Se supera l'altezza massima, vai alla pagina successiva
            if current_page_height > MAX_PAGE_HEIGHT:
                current_page += 1
                current_page_height = para_height  # Inizia nuova pagina con questo paragrafo
                page_map[i] = current_page
        
        return page_map
    
    def _paragraph_has_image(self, paragraph) -> bool:
        """Verifica se un paragrafo contiene un'immagine"""
        for run in paragraph.runs:
            run_element = run._element
            blips = run_element.findall('.//{*}blip')
            if blips:
                return True
        return False
    
    def _extract_image_label(self, text: str) -> Optional[str]:
        """
        Estrae label immagine da testo circostante.
        Supporta numerazione multi-livello (es: 5.2.2, 11.3.4.1) e diverse nomenclature.
        """
        # Pattern aggiornati per catturare numeri multi-livello
        # (\d+(?:\.\d+)*[a-z]?) cattura: 5, 5.2, 5.2.2, 5.2.2.1, 11a, etc.
        patterns = [
            (r'[Ee]sempio\.?\s*(\d+(?:\.\d+)*[a-z]?)', 'Esempio'),
            (r'[Ff]ig(?:ure|ura)?\.?\s*(\d+(?:\.\d+)*[a-z]?)', 'Fig'),
            (r'[Ii]mmagine\.?\s*(\d+(?:\.\d+)*[a-z]?)', 'Immagine'),
            (r'[Gg]rafico\.?\s*(\d+(?:\.\d+)*[a-z]?)', 'Grafico'),
            (r'[Dd]idascalia\.?\s*(\d+(?:\.\d+)*[a-z]?)', 'Didascalia'),
            (r'[Ss]chema\.?\s*(\d+(?:\.\d+)*[a-z]?)', 'Schema'),
            (r'[Dd]iagramma\.?\s*(\d+(?:\.\d+)*[a-z]?)', 'Diagramma'),
        ]
        
        for pattern, prefix in patterns:
            match = re.search(pattern, text, re.IGNORECASE)
            if match:
                number = match.group(1)
                return f"{prefix} {number}"
        
        return None
    
    def _extract_tables(self, page_map: Dict[int, int]) -> List[Dict]:
        """Estrae tabelle"""
        tables = []
        
        for i, table in enumerate(self.document.tables):
            rows_data = []
            for row in table.rows:
                row_data = [cell.text for cell in row.cells]
                rows_data.append(row_data)
            
            # Per le tabelle non abbiamo un paragraph_index diretto
            # Usiamo una stima basata sulla posizione nel documento
            # (in alternativa, si potrebbe cercare di trovare il paragrafo più vicino)
            estimated_para_index = i * 10  # Stima molto approssimativa
            page_number = page_map.get(estimated_para_index, 1)
            
            table_data = {
                'index': i,
                'rows': rows_data,
                'num_rows': len(rows_data),
                'num_cols': len(rows_data[0]) if rows_data else 0,
                'paragraph_index': None,
                'page_number': page_number  # NUOVO: numero di pagina stimato
            }
            
            tables.append(table_data)
        
        return tables
    
    def _extract_equations(self, page_map: Dict[int, int]) -> List[Dict]:
        """Estrae equazioni con classificazione automatica"""
        equations = []
        
        for i, paragraph in enumerate(self.document.paragraphs):
            omath_elements = paragraph._element.findall('.//{*}oMath')
            
            for omath in omath_elements:
                text = ''.join(omath.itertext())
                
                try:
                    from lxml import etree
                    mathml = etree.tostring(omath, encoding='unicode')
                except:
                    mathml = str(omath)
                
                context = self._get_context_5_lines(i)
                
                # CLASSIFICAZIONE AUTOMATICA
                category = 'generic'
                subcategory = None
                difficulty = 1
                tags = []
                
                if self.formula_classifier and text.strip():
                    try:
                        # Prova a classificare usando il testo matematico
                        classification = self.formula_classifier.classify(text)
                        category = classification.category
                        subcategory = classification.subcategory
                        difficulty = classification.difficulty
                        tags = classification.suggested_tags
                    except Exception as e:
                        # Se classificazione fallisce, usa valori default
                        print(f"Classificazione fallita per equazione: {e}")
                        category = 'other'
                        subcategory = 'misc'
                
                # Ottieni numero di pagina reale
                page_number = page_map.get(i, 1)
                
                eq_data = {
                    'paragraph_index': i,
                    'page_number': page_number,  # NUOVO: numero di pagina reale
                    'text_representation': text,
                    'mathml': mathml,
                    'context_before': context['before'],
                    'context_after': context['after'],
                    'category': category,
                    'subcategory': subcategory,
                    'difficulty': difficulty,
                    'tags': tags,
                    'sub_categories': []  # Mantenuto per compatibilità
                }
                
                equations.append(eq_data)
        
        return equations
    
    def _extract_references(self, page_map: Dict[int, int]) -> List[Dict]:
        """Estrae riferimenti"""
        references = []
        
        patterns = {
            'figure': r'[Ff]ig(?:ure|ura)?\.?\s*(\d+(?:\.\d+)?)',
            'table': r'[Tt]ab(?:ella)?\.?\s*(\d+(?:\.\d+)?)',
            'equation': r'[Ee]q(?:uation|uazione)?\.?\s*(\d+(?:\.\d+)?)',
        }
        
        for i, paragraph in enumerate(self.document.paragraphs):
            text = paragraph.text
            page_number = page_map.get(i, 1)
            
            for ref_type, pattern in patterns.items():
                for match in re.finditer(pattern, text):
                    references.append({
                        'type': ref_type,
                        'label': match.group(0),
                        'number': match.group(1),
                        'paragraph_index': i,
                        'page_number': page_number,  # NUOVO: numero di pagina reale
                        'position': match.start()
                    })
        
        return references
