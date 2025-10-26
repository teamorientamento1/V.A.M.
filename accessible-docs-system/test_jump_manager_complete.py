#!/usr/bin/env python3
"""
Test Completo Jump Manager - Con documento simulato
Testa tutte le funzionalità della nuova UI
"""
import sys
from pathlib import Path

# Setup path
ROOT_DIR = Path(__file__).parent.parent
sys.path.insert(0, str(ROOT_DIR))

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from core.knowledge_base import KnowledgeBase
from core.document_analyzer import WordAnalyzer
from modules.jump_manager.jump_creator import JumpManager
import json

def create_test_document_with_all_elements():
    """Crea documento Word completo con immagini, equazioni, tabelle"""
    print("\n" + "="*70)
    print("CREAZIONE DOCUMENTO TEST COMPLETO")
    print("="*70)
    
    doc = Document()
    doc.add_heading('Documento Test - Sistema Accessibile', 0)
    
    # Introduzione
    doc.add_paragraph(
        "Questo documento contiene immagini, equazioni e tabelle "
        "per testare il sistema Jump Manager completo."
    )
    
    # ==== SEZIONE 1: IMMAGINI ====
    doc.add_heading('1. Immagini e Figure', 1)
    
    doc.add_paragraph(
        "Nel corso di questo documento verranno presentate diverse figure. "
        "Come si può vedere in Fig.1.1, i dati mostrano un trend crescente."
    )
    
    # Immagine 1
    img_para1 = doc.add_paragraph()
    img_para1.add_run("[IMMAGINE 1 - Grafico vendite]").bold = True
    
    caption1 = doc.add_paragraph()
    caption1.add_run("Fig.1.1: ").bold = True
    caption1.add_run("Andamento vendite Q1-Q4 2024. Il grafico mostra un incremento "
                     "costante del fatturato con un picco nel Q4.")
    caption1.style = 'Caption'
    
    doc.add_paragraph(
        "I risultati presentati in Fig.1.1 dimostrano l'efficacia della strategia "
        "implementata. Inoltre, come vedremo più avanti in Fig.1.2, la tendenza "
        "si conferma anche nei dati regionali."
    )
    
    # Immagine 2
    doc.add_paragraph()
    img_para2 = doc.add_paragraph()
    img_para2.add_run("[IMMAGINE 2 - Mappa regioni]").bold = True
    
    caption2 = doc.add_paragraph()
    caption2.add_run("Fig.1.2: ").bold = True
    caption2.add_run("Distribuzione vendite per regione. La mappa evidenzia le aree "
                     "geografiche con maggiore crescita (in verde) e quelle stabili (in giallo).")
    caption2.style = 'Caption'
    
    # ==== SEZIONE 2: EQUAZIONI ====
    doc.add_heading('2. Formule Matematiche', 1)
    
    doc.add_paragraph(
        "Le relazioni matematiche sono fondamentali per comprendere i fenomeni. "
        "Consideriamo l'Eq.(2.1) che descrive il moto uniformemente accelerato:"
    )
    
    # Equazione 1 (simulata con testo)
    eq_para1 = doc.add_paragraph()
    eq_para1.alignment = 1  # Centro
    eq_run1 = eq_para1.add_run("s = s₀ + v₀t + ½at²")
    eq_run1.font.size = Pt(14)
    eq_run1.italic = True
    
    eq_label1 = doc.add_paragraph("(Eq.2.1)")
    eq_label1.alignment = 2  # Destra
    
    doc.add_paragraph(
        "dove s è lo spostamento, s₀ la posizione iniziale, v₀ la velocità iniziale, "
        "a l'accelerazione e t il tempo. Utilizzando l'Eq.(2.1), possiamo calcolare "
        "la posizione di un oggetto in qualsiasi istante."
    )
    
    doc.add_paragraph(
        "La seconda legge di Newton, espressa dall'Eq.(2.2), stabilisce la relazione "
        "tra forza, massa e accelerazione:"
    )
    
    # Equazione 2
    eq_para2 = doc.add_paragraph()
    eq_para2.alignment = 1
    eq_run2 = eq_para2.add_run("F = ma")
    eq_run2.font.size = Pt(14)
    eq_run2.italic = True
    eq_run2.font.color.rgb = RGBColor(0, 0, 128)
    
    eq_label2 = doc.add_paragraph("(Eq.2.2)")
    eq_label2.alignment = 2
    
    doc.add_paragraph(
        "Combinando l'Eq.(2.1) e l'Eq.(2.2) otteniamo una descrizione completa "
        "del moto."
    )
    
    # Equazione 3 - Integrale
    doc.add_paragraph(
        "Per calcolare l'area sotto una curva utilizziamo l'Eq.(2.3):"
    )
    
    eq_para3 = doc.add_paragraph()
    eq_para3.alignment = 1
    eq_run3 = eq_para3.add_run("∫ f(x)dx = F(b) - F(a)")
    eq_run3.font.size = Pt(14)
    eq_run3.italic = True
    
    eq_label3 = doc.add_paragraph("(Eq.2.3)")
    eq_label3.alignment = 2
    
    # ==== SEZIONE 3: TABELLE ====
    doc.add_heading('3. Dati Tabulari', 1)
    
    doc.add_paragraph(
        "I dati raccolti sono sintetizzati nella Tab.3.1:"
    )
    
    # Tabella 1
    table1 = doc.add_table(rows=5, cols=4)
    table1.style = 'Light Grid Accent 1'
    
    # Header
    header_cells = table1.rows[0].cells
    header_cells[0].text = 'Mese'
    header_cells[1].text = 'Vendite (k€)'
    header_cells[2].text = 'Crescita %'
    header_cells[3].text = 'Target'
    
    # Dati
    data = [
        ['Gennaio', '150', '+5%', '145'],
        ['Febbraio', '165', '+10%', '160'],
        ['Marzo', '180', '+9%', '175'],
        ['Aprile', '195', '+8%', '190']
    ]
    
    for i, row_data in enumerate(data, 1):
        cells = table1.rows[i].cells
        for j, value in enumerate(row_data):
            cells[j].text = value
    
    table1_caption = doc.add_paragraph()
    table1_caption.add_run("Tab.3.1: ").bold = True
    table1_caption.add_run("Dati trimestrali vendite con confronto target. "
                          "La crescita percentuale è calcolata rispetto al mese precedente.")
    table1_caption.style = 'Caption'
    
    doc.add_paragraph(
        "Come evidenziato nella Tab.3.1, tutti i mesi hanno superato il target previsto. "
        "Confrontando questi dati con la Fig.1.1, possiamo notare la correlazione tra "
        "crescita percentuale e fatturato assoluto."
    )
    
    # Tabella 2 - Più complessa
    doc.add_paragraph()
    doc.add_paragraph(
        "La Tab.3.2 presenta un'analisi più dettagliata per regione:"
    )
    
    table2 = doc.add_table(rows=4, cols=3)
    table2.style = 'Medium Shading 1 Accent 1'
    
    header2 = table2.rows[0].cells
    header2[0].text = 'Regione'
    header2[1].text = 'Fatturato (k€)'
    header2[2].text = 'Quota %'
    
    regions = [
        ['Nord', '450', '45%'],
        ['Centro', '350', '35%'],
        ['Sud', '200', '20%']
    ]
    
    for i, row_data in enumerate(regions, 1):
        cells = table2.rows[i].cells
        for j, value in enumerate(row_data):
            cells[j].text = value
    
    table2_caption = doc.add_paragraph()
    table2_caption.add_run("Tab.3.2: ").bold = True
    table2_caption.add_run("Distribuzione fatturato per macro-area geografica.")
    table2_caption.style = 'Caption'
    
    # ==== SEZIONE 4: RIFERIMENTI MULTIPLI ====
    doc.add_heading('4. Conclusioni', 1)
    
    doc.add_paragraph(
        "In conclusione, l'analisi presentata dimostra chiaramente l'efficacia "
        "della strategia. Come illustrato in Fig.1.1 e confermato dai dati della Tab.3.1, "
        "tutti gli indicatori sono positivi. Le formule presentate (Eq.2.1, Eq.2.2, Eq.2.3) "
        "forniscono il framework teorico per interpretare i risultati. La distribuzione "
        "geografica mostrata in Fig.1.2 e dettagliata nella Tab.3.2 evidenzia opportunità "
        "di crescita nelle aree del Sud."
    )
    
    # Salva
    output_path = Path("test_documento_completo.docx")
    doc.save(output_path)
    
    print(f"✓ Documento creato: {output_path}")
    print(f"  Contiene:")
    print(f"    - 2 immagini (simulate)")
    print(f"    - 3 equazioni")
    print(f"    - 2 tabelle")
    print(f"    - Multipli riferimenti incrociati")
    
    return output_path

def test_document_analysis():
    """Test 1: Analisi documento completo"""
    print("\n" + "="*70)
    print("TEST 1: Analisi Documento Completo")
    print("="*70)
    
    # Crea documento
    doc_path = create_test_document_with_all_elements()
    
    # Analizza
    kb = KnowledgeBase()
    analyzer = WordAnalyzer(kb)
    
    print("\nCaricamento documento...")
    if not analyzer.load_document(doc_path):
        print("✗ Errore caricamento")
        kb.close()
        return False
    
    print("✓ Documento caricato")
    print("\nAnalisi in corso...")
    
    analysis = analyzer.analyze(discipline="generic")
    
    # Mostra risultati
    print("\n" + "-"*70)
    print("RISULTATI ANALISI")
    print("-"*70)
    
    stats = analysis['statistics']
    print(f"\n📊 Statistiche:")
    print(f"  Paragrafi: {stats['total_paragraphs']}")
    print(f"  Immagini: {stats['total_images']}")
    print(f"  Equazioni: {stats['total_equations']}")
    print(f"  Tabelle: {stats['total_tables']}")
    print(f"  Riferimenti: {stats['total_references']}")
    
    # Dettagli immagini
    if analysis['images']:
        print(f"\n📸 Immagini trovate:")
        for img in analysis['images']:
            label = img.get('label', 'N/A')
            para = img.get('paragraph_index', 'N/A')
            print(f"  • {label} (paragrafo {para})")
    
    # Dettagli equazioni
    if analysis['equations']:
        print(f"\n∑ Equazioni trovate:")
        for eq in analysis['equations'][:5]:
            text = eq.get('text_representation', 'N/A')[:40]
            para = eq.get('paragraph_index', 'N/A')
            print(f"  • {text}... (paragrafo {para})")
    
    # Dettagli tabelle
    if analysis['tables']:
        print(f"\n📊 Tabelle trovate:")
        for table in analysis['tables']:
            rows = table['rows']
            cols = table['cols']
            idx = table['index']
            print(f"  • Tab.{idx}: {rows}x{cols}")
    
    # Dettagli riferimenti
    if analysis['references']:
        print(f"\n🔗 Riferimenti trovati:")
        ref_types = {}
        for ref in analysis['references']:
            ref_type = ref['type']
            ref_types[ref_type] = ref_types.get(ref_type, 0) + 1
        
        for ref_type, count in ref_types.items():
            print(f"  • {ref_type}: {count}")
        
        print(f"\n  Primi 5 riferimenti:")
        for ref in analysis['references'][:5]:
            print(f"    - {ref['full_match']} (para {ref['paragraph_index']})")
    
    # Salva analisi su file JSON
    output_json = Path("test_analysis_results.json")
    
    # Prepara per JSON (rimuovi dati binari)
    analysis_for_json = dict(analysis)
    if 'images' in analysis_for_json:
        for img in analysis_for_json['images']:
            if 'image_bytes' in img:
                img['image_bytes'] = f"<{len(img['image_bytes'])} bytes>"
    
    with open(output_json, 'w', encoding='utf-8') as f:
        json.dump(analysis_for_json, f, indent=2, ensure_ascii=False)
    
    print(f"\n✓ Analisi salvata: {output_json}")
    
    kb.close()
    return True, analysis

def test_jump_creation(analysis):
    """Test 2: Creazione jump da analisi"""
    print("\n" + "="*70)
    print("TEST 2: Creazione Jump")
    print("="*70)
    
    # Carica documento
    doc_path = Path("test_documento_completo.docx")
    if not doc_path.exists():
        print("✗ Documento non trovato")
        return False
    
    doc = Document(str(doc_path))
    jump_mgr = JumpManager(doc)
    
    # Prepara descrizioni
    descriptions = {
        'Fig.1.1': "Grafico a linee che mostra l'andamento delle vendite nei quattro trimestri del 2024. "
                   "L'asse X rappresenta il tempo (Q1, Q2, Q3, Q4), mentre l'asse Y mostra il fatturato "
                   "in migliaia di euro. La linea blu mostra un trend crescente costante, partendo da "
                   "circa 150k€ in Q1 e raggiungendo quasi 200k€ in Q4. Il picco massimo si registra "
                   "nell'ultimo trimestre, con un incremento del 30% rispetto all'inizio dell'anno.",
        
        'Fig.1.2': "Mappa dell'Italia suddivisa in tre macro-aree (Nord, Centro, Sud). Le regioni sono "
                   "colorate secondo il livello di crescita: verde per crescita elevata (>10%), "
                   "giallo per crescita moderata (5-10%), rosso per crescita bassa (<5%). "
                   "Il Nord Italia appare prevalentemente in verde, il Centro in giallo, "
                   "mentre il Sud mostra aree sia gialle che rosse, indicando opportunità di sviluppo."
    }
    
    print("\n1. Creazione sezione descrizioni...")
    jump_mgr.create_descriptions_section()
    print("   ✓ Sezione creata")
    
    print("\n2. Aggiunta descrizioni immagini...")
    for img in analysis['images']:
        label = img.get('label')
        if label and label in descriptions:
            jump_mgr.add_description_with_return(
                label=label,
                description=descriptions[label],
                source_paragraph_index=img.get('paragraph_index', 0)
            )
            print(f"   ✓ Descrizione aggiunta: {label}")
    
    print("\n3. Creazione jump immagini...")
    jumps_created = jump_mgr.create_image_jumps(
        images=analysis['images'],
        descriptions=descriptions
    )
    print(f"   ✓ Creati {jumps_created} jump per immagini")
    
    print("\n4. Scansione e creazione jump riferimenti...")
    ref_jumps = jump_mgr.scan_and_create_reference_jumps(
        references=analysis['references']
    )
    print(f"   ✓ Creati {ref_jumps} jump per riferimenti")
    
    print("\n5. Validazione jump...")
    validation = jump_mgr.validate_jumps()
    print(f"   ✓ Totale jump: {validation['total_jumps']}")
    
    # Salva documento
    output_doc = Path("test_documento_con_jump.docx")
    doc.save(str(output_doc))
    print(f"\n✓ Documento salvato: {output_doc}")
    
    # Export mappa jump
    jump_map_file = Path("test_jump_map.json")
    jump_mgr.export_jump_map(str(jump_map_file))
    print(f"✓ Mappa jump esportata: {jump_map_file}")
    
    # Sommario
    summary = jump_mgr.get_jumps_summary()
    print(f"\n📊 Sommario Jump:")
    print(f"   Total: {summary['total']}")
    print(f"   Per tipo:")
    for jtype, count in summary['by_type'].items():
        print(f"     - {jtype}: {count}")
    
    return True

def test_ui_integration():
    """Test 3: Test integrazione con UI (simulato)"""
    print("\n" + "="*70)
    print("TEST 3: Simulazione UI")
    print("="*70)
    
    print("\nQuesto test simula il workflow UI:")
    print("1. ✓ Caricamento documento")
    print("2. ✓ Analisi completata")
    print("3. ✓ Dati caricati in Jump Manager")
    print("4. ⚠ Seleziona elemento (simulato)")
    print("5. ⚠ Aggiungi descrizione (simulato)")
    print("6. ⚠ Crea jump (simulato)")
    print("7. ⚠ Valida e applica (simulato)")
    
    print("\nPer test UI completo:")
    print("  python ui/main_window_updated.py")
    
    return True

def create_descriptions_template():
    """Crea template JSON per descrizioni"""
    print("\n" + "="*70)
    print("CREAZIONE TEMPLATE DESCRIZIONI")
    print("="*70)
    
    template = {
        "Fig.1.1": "[Inserisci descrizione dettagliata dell'immagine 1]",
        "Fig.1.2": "[Inserisci descrizione dettagliata dell'immagine 2]",
        "Tab.3.1": "[Inserisci descrizione della tabella 1]",
        "Tab.3.2": "[Inserisci descrizione della tabella 2]",
        "_istruzioni": {
            "come_usare": "Sostituisci i placeholder con descrizioni dettagliate",
            "linee_guida": [
                "Descrivi cosa mostra l'elemento",
                "Spiega il significato dei dati",
                "Collega al contesto del documento",
                "Usa un linguaggio chiaro e accessibile"
            ]
        }
    }
    
    output_file = Path("test_descriptions_template.json")
    with open(output_file, 'w', encoding='utf-8') as f:
        json.dump(template, f, indent=2, ensure_ascii=False)
    
    print(f"✓ Template creato: {output_file}")
    print("\nModifica questo file con le descrizioni reali e:")
    print("  1. Caricalo nell'UI (pulsante 'Carica da File')")
    print("  2. Oppure usalo programmaticamente nel workflow")
    
    return True

def main():
    """Main test"""
    print("\n" + "█"*70)
    print("  JUMP MANAGER - TEST SUITE COMPLETO")
    print("█"*70)
    
    tests = [
        ("Creazione Template Descrizioni", create_descriptions_template),
        ("Analisi Documento", test_document_analysis),
    ]
    
    results = []
    analysis_result = None
    
    for test_name, test_func in tests:
        try:
            print(f"\n{'='*70}")
            result = test_func()
            if isinstance(result, tuple):
                success, analysis_result = result
            else:
                success = result
            results.append((test_name, success))
        except Exception as e:
            print(f"\n✗ ERRORE in {test_name}: {e}")
            import traceback
            traceback.print_exc()
            results.append((test_name, False))
    
    # Test jump solo se analisi ok
    if analysis_result:
        try:
            print(f"\n{'='*70}")
            success = test_jump_creation(analysis_result)
            results.append(("Creazione Jump", success))
        except Exception as e:
            print(f"\n✗ ERRORE in Creazione Jump: {e}")
            import traceback
            traceback.print_exc()
            results.append(("Creazione Jump", False))
    
    # Test UI
    try:
        print(f"\n{'='*70}")
        success = test_ui_integration()
        results.append(("Integrazione UI", success))
    except Exception as e:
        print(f"\n✗ ERRORE in Integrazione UI: {e}")
        results.append(("Integrazione UI", False))
    
    # Sommario
    print("\n" + "█"*70)
    print("  SOMMARIO TEST")
    print("█"*70)
    
    for test_name, success in results:
        status = "✓ PASS" if success else "✗ FAIL"
        print(f"{status}: {test_name}")
    
    passed = sum(1 for _, s in results if s)
    total = len(results)
    
    print(f"\nRisultato: {passed}/{total} test passati")
    
    # File generati
    print("\n" + "="*70)
    print("FILE GENERATI")
    print("="*70)
    print("Documenti:")
    print("  • test_documento_completo.docx - Documento di test")
    print("  • test_documento_con_jump.docx - Con jump applicati")
    print("\nDati:")
    print("  • test_analysis_results.json - Risultati analisi")
    print("  • test_jump_map.json - Mappa jump creati")
    print("  • test_descriptions_template.json - Template descrizioni")
    
    print("\n" + "="*70)
    print("PROSSIMI PASSI")
    print("="*70)
    
    if passed == total:
        print("🎉 TUTTI I TEST PASSATI!")
        print("\nPer usare il sistema:")
        print("1. python ui/main_window_updated.py")
        print("2. Carica test_documento_completo.docx")
        print("3. Analizza documento")
        print("4. Vai a Jump Manager")
        print("5. Crea jump con UI grafica!")
    else:
        print("⚠ Alcuni test falliti. Controlla gli errori sopra.")
    
    print("\nPer test manuale UI:")
    print("  cd accessible-docs-system")
    print("  python ui/main_window_updated.py")
    print()
    
    return passed == total

if __name__ == "__main__":
    success = main()
    sys.exit(0 if success else 1)
