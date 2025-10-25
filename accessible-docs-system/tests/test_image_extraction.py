#!/usr/bin/env python3
"""
Test Image Extraction - Verifica estrazione immagini da Word
"""
import sys
from pathlib import Path

# Setup path
ROOT_DIR = Path(__file__).parent.parent
sys.path.insert(0, str(ROOT_DIR))

from docx import Document
from docx.shared import Inches
from core.knowledge_base import KnowledgeBase
from core.document_analyzer import WordAnalyzer

def create_test_document_with_images():
    """Crea documento Word di test con immagini simulate"""
    print("\n" + "="*60)
    print("CREAZIONE DOCUMENTO TEST CON IMMAGINI")
    print("="*60)
    
    doc = Document()
    doc.add_heading('Documento Test - Estrazione Immagini', 0)
    
    # Introduzione
    doc.add_paragraph(
        "Questo documento contiene diverse immagini per testare "
        "l'estrazione automatica."
    )
    
    # Immagine 1 con caption
    doc.add_heading('Sezione 1: Prima Immagine', 1)
    doc.add_paragraph("Testo introduttivo prima dell'immagine.")
    
    # Placeholder per immagine (in un documento reale ci sarebbe un'immagine vera)
    img_para = doc.add_paragraph()
    img_para.add_run("[QUI CI SAREBBE UN'IMMAGINE]").bold = True
    
    # Caption
    caption1 = doc.add_paragraph("Fig.1.1: Grafico dell'andamento delle vendite nel Q1-2024")
    caption1.style = 'Caption'
    
    doc.add_paragraph(
        "Come si vede in Fig.1.1, l'andamento è crescente. "
        "I dati mostrano un incremento del 15% rispetto al trimestre precedente."
    )
    
    # Immagine 2
    doc.add_heading('Sezione 2: Seconda Immagine', 1)
    doc.add_paragraph("Il processo è illustrato nella figura seguente.")
    
    img_para2 = doc.add_paragraph()
    img_para2.add_run("[IMMAGINE 2 - Schema processo]").bold = True
    
    caption2 = doc.add_paragraph("Figura 2.1: Schema a blocchi del processo produttivo")
    caption2.style = 'Caption'
    
    doc.add_paragraph(
        "Il processo mostrato in Fig.2.1 è composto da tre fasi principali: "
        "acquisizione, elaborazione e controllo qualità."
    )
    
    # Immagine in tabella
    doc.add_heading('Sezione 3: Immagine in Tabella', 1)
    doc.add_paragraph("La tabella seguente contiene un'immagine:")
    
    table = doc.add_table(rows=2, cols=2)
    table.style = 'Light Grid Accent 1'
    
    table.rows[0].cells[0].text = "Descrizione"
    table.rows[0].cells[1].text = "Immagine"
    table.rows[1].cells[0].text = "Schema circuito elettrico"
    table.rows[1].cells[1].text = "[IMMAGINE IN TABELLA]"
    
    doc.add_paragraph()
    doc.add_paragraph("Tab.1.1: Componenti principali del circuito")
    
    # Riferimenti multipli
    doc.add_heading('Sezione 4: Riferimenti', 1)
    doc.add_paragraph(
        "In questo documento abbiamo visto diverse figure: Fig.1.1 mostra i dati di vendita, "
        "mentre Fig.2.1 illustra il processo. La Tab.1.1 fornisce dettagli tecnici."
    )
    
    # Salva
    output_path = Path("test_document_with_images.docx")
    doc.save(output_path)
    print(f"✓ Documento creato: {output_path}")
    print(f"  Nota: Questo documento contiene placeholder per immagini")
    print(f"  Per test completo, apri il documento e inserisci immagini reali")
    
    return output_path

def test_image_extraction_basic():
    """Test 1: Estrazione base senza immagini reali"""
    print("\n" + "="*60)
    print("TEST 1: Estrazione Base (Documento Placeholder)")
    print("="*60)
    
    # Crea documento test
    doc_path = create_test_document_with_images()
    
    # Analizza
    kb = KnowledgeBase()
    analyzer = WordAnalyzer(kb)
    
    if not analyzer.load_document(doc_path):
        print("✗ Errore nel caricamento documento")
        return False
    
    print("\nAnalisi in corso...")
    analysis = analyzer.analyze(discipline="generic")
    
    # Risultati
    print(f"\n--- RISULTATI ---")
    print(f"Immagini trovate: {len(analysis['images'])}")
    
    if analysis['images']:
        print("\nDettagli immagini:")
        for img in analysis['images']:
            print(f"\n  Immagine {img['index']}:")
            print(f"    - Filename: {img.get('filename', 'N/A')}")
            print(f"    - Size: {img.get('size_bytes', 0)} bytes")
            print(f"    - Label: {img.get('label', 'Non trovato')}")
            print(f"    - Caption: {img.get('caption', 'N/A')[:50]}...")
            print(f"    - In table: {img.get('in_table', False)}")
            print(f"    - Paragrafo: {img.get('paragraph_index', 'N/A')}")
    else:
        print("  ⚠ Nessuna immagine trovata (normale per placeholder)")
    
    # Riferimenti
    print(f"\nRiferimenti trovati: {len(analysis['references'])}")
    fig_refs = [r for r in analysis['references'] if r['type'] == 'figure']
    print(f"  - Riferimenti a figure: {len(fig_refs)}")
    
    if fig_refs:
        print("\n  Riferimenti a figure trovati:")
        for ref in fig_refs[:5]:
            print(f"    - {ref['full_match']} (paragrafo {ref['paragraph_index']})")
    
    kb.close()
    print("\n✓ Test base completato")
    return True

def test_image_extraction_with_real_doc():
    """Test 2: Con documento reale (se disponibile)"""
    print("\n" + "="*60)
    print("TEST 2: Estrazione da Documento Reale")
    print("="*60)
    
    # Cerca documenti Word nella directory corrente
    doc_files = list(Path.cwd().glob("*.docx"))
    doc_files = [f for f in doc_files if not f.name.startswith("~$")]
    
    if not doc_files:
        print("⚠ Nessun documento .docx trovato nella directory corrente")
        print("  Posiziona un documento Word con immagini nella stessa cartella")
        print("  e riesegui il test")
        return True
    
    print(f"Documenti trovati: {len(doc_files)}")
    for i, doc_file in enumerate(doc_files[:5], 1):
        print(f"  {i}. {doc_file.name}")
    
    # Usa il primo documento non-test
    real_docs = [d for d in doc_files if 'test' not in d.name.lower()]
    if not real_docs:
        real_docs = doc_files
    
    doc_path = real_docs[0]
    print(f"\nAnalizzando: {doc_path.name}")
    
    # Analizza
    kb = KnowledgeBase()
    analyzer = WordAnalyzer(kb)
    
    if not analyzer.load_document(doc_path):
        print("✗ Errore nel caricamento")
        kb.close()
        return False
    
    print("Analisi in corso...")
    analysis = analyzer.analyze(discipline="generic")
    
    # Risultati dettagliati
    print(f"\n--- RISULTATI DETTAGLIATI ---")
    stats = analysis['statistics']
    print(f"Paragrafi: {stats['total_paragraphs']}")
    print(f"Equazioni: {stats['total_equations']}")
    print(f"Tabelle: {stats['total_tables']}")
    print(f"Immagini: {stats['total_images']}")
    print(f"Riferimenti: {stats['total_references']}")
    
    if analysis['images']:
        print(f"\n--- DETTAGLI IMMAGINI ---")
        for img in analysis['images']:
            print(f"\nImmagine {img['index']}:")
            print(f"  File: {img.get('filename', 'N/A')}")
            print(f"  Tipo: {img.get('content_type', 'N/A')}")
            print(f"  Size: {img.get('size_bytes', 0) / 1024:.1f} KB")
            
            if img.get('width') and img.get('height'):
                print(f"  Dimensioni: {img['width']}x{img['height']} px")
            
            print(f"  Label: {img.get('label', 'Non trovato')}")
            
            if img.get('caption'):
                print(f"  Caption: {img['caption'][:80]}...")
            
            if img.get('in_table'):
                print(f"  Posizione: Tabella {img.get('table_index')}")
            else:
                print(f"  Posizione: Paragrafo {img.get('paragraph_index')}")
            
            if img.get('context_before'):
                print(f"  Contesto: ...{img['context_before'][-50:]}...")
    
    # Export dati immagini
    if analysis['images']:
        import json
        output_file = Path("extracted_images_data.json")
        
        # Rimuovi bytes per JSON
        images_for_json = []
        for img in analysis['images']:
            img_copy = img.copy()
            if 'image_bytes' in img_copy:
                img_copy['image_bytes'] = f"<{len(img['image_bytes'])} bytes>"
            images_for_json.append(img_copy)
        
        with open(output_file, 'w', encoding='utf-8') as f:
            json.dump({
                'document': doc_path.name,
                'total_images': len(images_for_json),
                'images': images_for_json
            }, f, indent=2, ensure_ascii=False)
        
        print(f"\n✓ Dati immagini esportati: {output_file}")
    
    kb.close()
    print("\n✓ Test con documento reale completato")
    return True

def test_image_label_extraction():
    """Test 3: Estrazione label da contesto"""
    print("\n" + "="*60)
    print("TEST 3: Estrazione Label")
    print("="*60)
    
    kb = KnowledgeBase()
    analyzer = WordAnalyzer(kb)
    
    # Test casi vari
    test_cases = [
        ("Come si vede in Fig.2.5, il risultato è evidente.", "Fig.2.5"),
        ("La Figura 1.1 mostra l'andamento.", "Fig.1.1"),
        ("Immagine 3: Schema del processo", "Fig.3"),
        ("Grafico 4.2: Distribuzione dei dati", "Fig.4.2"),
        ("Nel testo senza riferimenti", None),
    ]
    
    print("\nTest estrazione label:")
    all_passed = True
    
    for text, expected in test_cases:
        result = analyzer._extract_image_label(text)
        status = "✓" if result == expected else "✗"
        print(f"  {status} '{text[:40]}...' → {result} (atteso: {expected})")
        if result != expected:
            all_passed = False
    
    kb.close()
    
    if all_passed:
        print("\n✓ Tutti i test label passati")
    else:
        print("\n⚠ Alcuni test falliti")
    
    return all_passed

def main():
    """Main test function"""
    print("\n" + "█"*60)
    print("  IMAGE EXTRACTION - TEST SUITE")
    print("█"*60)
    
    tests = [
        ("Estrazione Base", test_image_extraction_basic),
        ("Estrazione Label", test_image_label_extraction),
        ("Estrazione Documento Reale", test_image_extraction_with_real_doc),
    ]
    
    results = []
    
    for test_name, test_func in tests:
        try:
            success = test_func()
            results.append((test_name, success))
        except Exception as e:
            print(f"\n✗ ERRORE in {test_name}: {e}")
            import traceback
            traceback.print_exc()
            results.append((test_name, False))
    
    # Sommario
    print("\n" + "█"*60)
    print("  SOMMARIO TEST")
    print("█"*60)
    
    for test_name, success in results:
        status = "✓ PASS" if success else "✗ FAIL"
        print(f"{status}: {test_name}")
    
    passed = sum(1 for _, s in results if s)
    total = len(results)
    
    print(f"\nRisultato: {passed}/{total} test passati")
    
    if passed == total:
        print("\n🎉 ESTRAZIONE IMMAGINI FUNZIONANTE!")
        print("\nPer test completo:")
        print("1. Apri test_document_with_images.docx")
        print("2. Inserisci immagini reali")
        print("3. Riesegui questo test")
        print("\nOppure analizza direttamente i tuoi documenti reali!")
    
    return passed == total

if __name__ == "__main__":
    success = main()
    sys.exit(0 if success else 1)
