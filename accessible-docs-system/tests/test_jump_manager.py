#!/usr/bin/env python3
"""
Test Jump Manager - Verifica funzionalità collegamenti ipertestuali
"""
import sys
from pathlib import Path

# Setup path
ROOT_DIR = Path(__file__).parent.parent
sys.path.insert(0, str(ROOT_DIR))

from docx import Document
from modules.jump_manager.jump_creator import JumpManager

def test_basic_hyperlink_bookmark():
    """Test 1: Creazione base hyperlink e bookmark"""
    print("\n" + "="*60)
    print("TEST 1: Hyperlink e Bookmark Base")
    print("="*60)
    
    # Crea documento test
    doc = Document()
    doc.add_heading('Test Jump Manager', 0)
    
    # Paragrafo 1 con bookmark
    para1 = doc.add_paragraph("Questo è il paragrafo TARGET. ")
    para1.add_run("Qui arriva il link!").bold = True
    
    # Crea Jump Manager
    jump_mgr = JumpManager(doc)
    
    # Crea bookmark
    jump_mgr.create_bookmark(para1, "test_bookmark_1")
    print("✓ Bookmark 'test_bookmark_1' creato")
    
    # Paragrafo 2 con link
    para2 = doc.add_paragraph("Clicca qui per andare al target: ")
    jump_mgr.create_hyperlink(para2, "VAI AL TARGET", "test_bookmark_1")
    print("✓ Hyperlink creato")
    
    # Salva
    output_path = Path("test_jump_basic.docx")
    doc.save(output_path)
    print(f"\n✓ Documento salvato: {output_path}")
    print("  Apri in Word e testa il link!")
    
    return True

def test_multiple_jumps():
    """Test 2: Multipli jump e bookmark"""
    print("\n" + "="*60)
    print("TEST 2: Multipli Jump")
    print("="*60)
    
    doc = Document()
    doc.add_heading('Test Multipli Jump', 0)
    jump_mgr = JumpManager(doc)
    
    # Crea 3 sezioni con bookmark
    sections = []
    for i in range(1, 4):
        section_para = doc.add_paragraph()
        heading = section_para.add_run(f"SEZIONE {i}")
        heading.bold = True
        heading.font.size = 14
        
        bookmark_name = f"section_{i}"
        jump_mgr.create_bookmark(section_para, bookmark_name)
        sections.append(bookmark_name)
        
        doc.add_paragraph(f"Contenuto della sezione {i}...")
        doc.add_paragraph()  # Spazio
        print(f"✓ Sezione {i} con bookmark creata")
    
    # Crea indice con link
    doc.add_page_break()
    doc.add_heading('INDICE', 1)
    
    index_para = doc.add_paragraph()
    for i, bookmark in enumerate(sections, 1):
        index_para.add_run(f"{i}. ")
        jump_mgr.create_hyperlink(index_para, f"Sezione {i}", bookmark)
        index_para.add_run("\n")
    
    print("✓ Indice con link creato")
    
    # Salva
    output_path = Path("test_jump_multiple.docx")
    doc.save(output_path)
    print(f"\n✓ Documento salvato: {output_path}")
    print("  Testa i 3 link dall'indice!")
    
    return True

def test_image_jumps_simulation():
    """Test 3: Simulazione jump per immagini (senza immagini vere)"""
    print("\n" + "="*60)
    print("TEST 3: Simulazione Jump Immagini")
    print("="*60)
    
    doc = Document()
    doc.add_heading('Test Jump Immagini (Simulato)', 0)
    jump_mgr = JumpManager(doc)
    
    # Simula documento con "immagini" (paragrafi placeholder)
    doc.add_paragraph("Testo introduttivo del documento...")
    doc.add_paragraph()
    
    # Immagine 1 (placeholder)
    img1_para = doc.add_paragraph()
    img1_para.add_run("[IMMAGINE 1 - Grafico andamento]").italic = True
    jump_mgr.create_hyperlink(img1_para, "[Vedi descrizione]", "DESC_Fig_2_1")
    doc.add_paragraph()
    
    doc.add_paragraph("Come si vede in Fig.2.1, l'andamento è crescente...")
    doc.add_paragraph()
    
    # Immagine 2 (placeholder)
    img2_para = doc.add_paragraph()
    img2_para.add_run("[IMMAGINE 2 - Schema processo]").italic = True
    jump_mgr.create_hyperlink(img2_para, "[Vedi descrizione]", "DESC_Fig_2_2")
    doc.add_paragraph()
    
    doc.add_paragraph("Il processo mostrato in Fig.2.2 è composto da 3 fasi...")
    
    # Crea sezione descrizioni
    jump_mgr.create_descriptions_section()
    
    # Aggiungi descrizioni
    jump_mgr.add_description_with_return(
        label="Fig.2.1",
        description="Grafico che mostra l'andamento delle vendite nel trimestre Q1-2024. "
                   "L'asse X rappresenta i mesi (gennaio, febbraio, marzo), mentre l'asse Y "
                   "mostra il fatturato in migliaia di euro. Si nota un trend crescente costante.",
        source_paragraph_index=2
    )
    
    jump_mgr.add_description_with_return(
        label="Fig.2.2",
        description="Schema a blocchi del processo produttivo. Il diagramma illustra le tre "
                   "fasi principali: (1) Acquisizione materie prime, (2) Lavorazione, "
                   "(3) Controllo qualità. Le frecce indicano il flusso del processo.",
        source_paragraph_index=6
    )
    
    print("✓ 2 immagini simulate con jump creati")
    print("✓ Sezione descrizioni creata")
    
    # Statistiche
    summary = jump_mgr.get_jumps_summary()
    print(f"\nSommario jump:")
    print(f"  - Totale: {summary['total']}")
    
    # Salva
    output_path = Path("test_jump_images.docx")
    doc.save(output_path)
    print(f"\n✓ Documento salvato: {output_path}")
    print("  Apri e testa i jump alle descrizioni!")
    
    return True

def test_reference_jumps():
    """Test 4: Jump da riferimenti nel testo"""
    print("\n" + "="*60)
    print("TEST 4: Jump da Riferimenti")
    print("="*60)
    
    doc = Document()
    doc.add_heading('Test Riferimenti', 0)
    jump_mgr = JumpManager(doc)
    
    # Simula riferimenti come quelli trovati dall'analyzer
    references = [
        {
            'type': 'figure',
            'label': '2.1',
            'full_match': 'Fig.2.1',
            'paragraph_index': 1,
            'position_in_paragraph': 20
        },
        {
            'type': 'equation',
            'label': '3',
            'full_match': 'Eq.(3)',
            'paragraph_index': 2,
            'position_in_paragraph': 15
        }
    ]
    
    # Crea paragrafi con riferimenti
    doc.add_paragraph("Testo introduttivo...")
    doc.add_paragraph("Come mostrato in Fig.2.1, i risultati sono significativi.")
    doc.add_paragraph("Usando l'Eq.(3) otteniamo il valore finale.")
    
    # Converti riferimenti in jump
    jumps_created = jump_mgr.scan_and_create_reference_jumps(references)
    
    print(f"✓ {jumps_created} jump da riferimenti creati")
    
    # Salva
    output_path = Path("test_jump_references.docx")
    doc.save(output_path)
    print(f"\n✓ Documento salvato: {output_path}")
    
    return True

def main():
    """Main test function"""
    print("\n" + "█"*60)
    print("  JUMP MANAGER - TEST SUITE")
    print("█"*60)
    
    tests = [
        ("Test Base Hyperlink/Bookmark", test_basic_hyperlink_bookmark),
        ("Test Multipli Jump", test_multiple_jumps),
        ("Test Jump Immagini (Simulato)", test_image_jumps_simulation),
        ("Test Jump Riferimenti", test_reference_jumps),
    ]
    
    results = []
    
    for test_name, test_func in tests:
        try:
            success = test_func()
            results.append((test_name, success))
        except Exception as e:
            print(f"\n✗ ERRORE in {test_name}: {e}")
            import traceback
            traceback.print_exc()
            results.append((test_name, False))
    
    # Sommario
    print("\n" + "█"*60)
    print("  SOMMARIO TEST")
    print("█"*60)
    
    for test_name, success in results:
        status = "✓ PASS" if success else "✗ FAIL"
        print(f"{status}: {test_name}")
    
    passed = sum(1 for _, s in results if s)
    total = len(results)
    
    print(f"\nRisultato: {passed}/{total} test passati")
    
    if passed == total:
        print("\n🎉 TUTTI I TEST COMPLETATI!")
        print("\nFile generati:")
        print("  - test_jump_basic.docx")
        print("  - test_jump_multiple.docx")
        print("  - test_jump_images.docx")
        print("  - test_jump_references.docx")
        print("\nApri questi file in Word e testa i link!")
    
    return passed == total

if __name__ == "__main__":
    success = main()
    sys.exit(0 if success else 1)
