#!/usr/bin/env python3
"""
Setup Rapido Jump Manager
Prepara tutto per iniziare subito a usare il sistema
"""
import sys
from pathlib import Path
import subprocess

def print_header(text):
    """Stampa header"""
    print("\n" + "="*70)
    print(f"  {text}")
    print("="*70)

def check_python_version():
    """Verifica versione Python"""
    print_header("CHECK PYTHON")
    version = sys.version_info
    print(f"Python {version.major}.{version.minor}.{version.micro}")
    
    if version.major < 3 or (version.major == 3 and version.minor < 8):
        print("✗ ERRORE: Serve Python 3.8 o superiore")
        return False
    
    print("✓ Versione Python OK")
    return True

def install_dependencies():
    """Installa dipendenze"""
    print_header("INSTALLAZIONE DIPENDENZE")
    
    requirements = [
        "python-docx",
        "pillow",
        "numpy",
    ]
    
    print("Installazione pacchetti base...")
    for package in requirements:
        try:
            print(f"  Installing {package}...", end=" ")
            subprocess.check_call(
                [sys.executable, "-m", "pip", "install", package],
                stdout=subprocess.DEVNULL,
                stderr=subprocess.DEVNULL
            )
            print("✓")
        except:
            print("✗ (prova manualmente)")
    
    print("\n✓ Dipendenze installate")
    return True

def create_directories():
    """Crea directory necessarie"""
    print_header("CREAZIONE DIRECTORY")
    
    dirs = [
        Path("data"),
        Path("data/knowledge_base"),
        Path("data/backups"),
        Path("data/user_data"),
        Path("logs"),
        Path("output"),
        Path("temp"),
    ]
    
    for directory in dirs:
        directory.mkdir(parents=True, exist_ok=True)
        print(f"  ✓ {directory}")
    
    print("\n✓ Directory create")
    return True

def test_imports():
    """Test import moduli"""
    print_header("TEST IMPORT")
    
    modules = [
        ("tkinter", "GUI"),
        ("docx", "python-docx"),
        ("PIL", "Pillow"),
    ]
    
    all_ok = True
    for module, name in modules:
        try:
            __import__(module)
            print(f"  ✓ {name}")
        except ImportError:
            print(f"  ✗ {name} - MANCANTE")
            all_ok = False
    
    return all_ok

def create_test_document():
    """Crea documento di test"""
    print_header("CREAZIONE DOCUMENTO TEST")
    
    try:
        from docx import Document
        from docx.shared import Pt
        
        doc = Document()
        doc.add_heading('Test Jump Manager', 0)
        
        # Intro
        doc.add_paragraph(
            "Questo è un documento di test per il Jump Manager. "
            "Contiene elementi dimostrativi."
        )
        
        # Sezione immagini
        doc.add_heading('1. Immagini', 1)
        doc.add_paragraph("Vedi Fig.1.1 per il grafico principale.")
        
        img_para = doc.add_paragraph()
        img_para.add_run("[IMMAGINE PLACEHOLDER]").bold = True
        
        caption = doc.add_paragraph()
        caption.add_run("Fig.1.1: ").bold = True
        caption.add_run("Grafico di esempio per test sistema.")
        
        # Sezione equazioni
        doc.add_heading('2. Equazioni', 1)
        doc.add_paragraph("La formula Eq.(2.1) definisce la relazione:")
        
        eq_para = doc.add_paragraph()
        eq_para.alignment = 1
        eq_run = eq_para.add_run("F = ma")
        eq_run.font.size = Pt(14)
        eq_run.italic = True
        
        eq_label = doc.add_paragraph("(Eq.2.1)")
        eq_label.alignment = 2
        
        # Sezione tabelle
        doc.add_heading('3. Tabelle', 1)
        doc.add_paragraph("I dati in Tab.3.1 mostrano i risultati:")
        
        table = doc.add_table(rows=3, cols=3)
        table.style = 'Light Grid Accent 1'
        
        # Header
        header = table.rows[0].cells
        header[0].text = 'Categoria'
        header[1].text = 'Valore'
        header[2].text = 'Percentuale'
        
        # Dati
        table.rows[1].cells[0].text = 'A'
        table.rows[1].cells[1].text = '100'
        table.rows[1].cells[2].text = '50%'
        
        table.rows[2].cells[0].text = 'B'
        table.rows[2].cells[1].text = '100'
        table.rows[2].cells[2].text = '50%'
        
        table_caption = doc.add_paragraph()
        table_caption.add_run("Tab.3.1: ").bold = True
        table_caption.add_run("Distribuzione dati campione.")
        
        # Conclusioni
        doc.add_heading('4. Conclusioni', 1)
        doc.add_paragraph(
            "Come mostrato in Fig.1.1 e confermato dalla Tab.3.1, "
            "utilizzando l'Eq.(2.1) otteniamo i risultati attesi."
        )
        
        # Salva
        output = Path("quickstart_test_document.docx")
        doc.save(output)
        
        print(f"✓ Documento creato: {output}")
        return True
        
    except Exception as e:
        print(f"✗ Errore: {e}")
        return False

def run_basic_test():
    """Test funzionamento base"""
    print_header("TEST FUNZIONAMENTO BASE")
    
    try:
        # Test Knowledge Base
        print("Test Knowledge Base...", end=" ")
        sys.path.insert(0, str(Path.cwd()))
        from core.knowledge_base import KnowledgeBase
        
        kb = KnowledgeBase()
        stats = kb.get_statistics()
        kb.close()
        print("✓")
        
        # Test Analyzer
        print("Test Document Analyzer...", end=" ")
        from core.document_analyzer import WordAnalyzer
        analyzer = WordAnalyzer(KnowledgeBase())
        print("✓")
        
        # Test Jump Manager
        print("Test Jump Manager...", end=" ")
        from modules.jump_manager.jump_creator import JumpManager
        from docx import Document
        doc = Document()
        jm = JumpManager(doc)
        print("✓")
        
        print("\n✓ Tutti i componenti funzionanti")
        return True
        
    except Exception as e:
        print(f"\n✗ Errore: {e}")
        import traceback
        traceback.print_exc()
        return False

def show_next_steps():
    """Mostra prossimi passi"""
    print_header("SETUP COMPLETATO!")
    
    print("""
🎉 Il sistema è pronto all'uso!

PROSSIMI PASSI:

1. AVVIA L'APPLICAZIONE
   python ui/main_window_updated.py

2. USA IL DOCUMENTO TEST
   - File: quickstart_test_document.docx
   - Caricalo nell'app
   - Vai a Jump Manager
   - Testa le funzionalità

3. LEGGI LA GUIDA
   - docs/JUMP_MANAGER_GUIDE.md
   - docs/QUICKSTART.md

4. TEST AVANZATI
   python tests/test_jump_manager_complete.py

COMANDI RAPIDI:

# Avvia UI
python ui/main_window_updated.py

# Test completo
python tests/test_jump_manager_complete.py

# Test sistema
python tests/test_system.py

RISORSE:

📖 Guide:
   - docs/JUMP_MANAGER_GUIDE.md (uso Jump Manager)
   - docs/DEVELOPMENT.md (sviluppo)
   - README.md (panoramica)

🧪 Test:
   - tests/test_jump_manager_complete.py
   - tests/test_system.py

💡 Esempi:
   - examples/workflow_completo.py
   - examples/esempio_uso.py

SUPPORTO:

❓ Problemi? Controlla:
   1. logs/system.log
   2. Riesegui: python setup_jump_manager.py
   3. Verifica: pip list | grep docx

📧 Per aiuto:
   - Leggi documentazione
   - Controlla esempi
   - Esegui test

Buon lavoro! 🚀
    """)

def main():
    """Main setup"""
    print("\n" + "█"*70)
    print("  JUMP MANAGER - SETUP RAPIDO")
    print("█"*70)
    print("\nQuesto script prepara il sistema per l'uso immediato.")
    print("Ci vorranno circa 2-3 minuti.\n")
    
    input("Premi ENTER per iniziare...")
    
    # Step 1: Check Python
    if not check_python_version():
        print("\n✗ Setup fallito: versione Python non supportata")
        return False
    
    # Step 2: Installa dipendenze
    if not install_dependencies():
        print("\n⚠ Warning: alcune dipendenze potrebbero mancare")
    
    # Step 3: Crea directory
    if not create_directories():
        print("\n✗ Setup fallito: errore creazione directory")
        return False
    
    # Step 4: Test import
    if not test_imports():
        print("\n⚠ Warning: alcuni moduli mancanti")
        print("   Installa manualmente: pip install python-docx pillow")
    
    # Step 5: Documento test
    if not create_test_document():
        print("\n⚠ Warning: documento test non creato")
        print("   Puoi crearlo manualmente dopo")
    
    # Step 6: Test funzionamento
    if not run_basic_test():
        print("\n⚠ Warning: alcuni test falliti")
        print("   Controlla che tutti i file siano presenti")
    
    # Step 7: Next steps
    show_next_steps()
    
    return True

if __name__ == "__main__":
    try:
        success = main()
        sys.exit(0 if success else 1)
    except KeyboardInterrupt:
        print("\n\nSetup interrotto dall'utente.")
        sys.exit(1)
    except Exception as e:
        print(f"\n\n✗ ERRORE FATALE: {e}")
        import traceback
        traceback.print_exc()
        sys.exit(1)
