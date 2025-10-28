"""
Jump Manager - Sistema di collegamenti con immagini e descrizioni visibili
"""

from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
import re


class JumpManager:
    """Gestione collegamenti ipertestuali con visualizzazione immagini e descrizioni"""
    
    def __init__(self, document):
        self.document = document
        self.jump_map = {}
        self.descriptions_start_index = None
        
    def create_image_jump_with_preview(self, image_data, description, return_text="↑ Torna al testo"):
        """
        Crea jump per immagine con link visibile e descrizione formattata
        
        Args:
            image_data: Dict con 'paragraph_index', 'label', 'image_element'
            description: Testo della descrizione
            return_text: Testo del link di ritorno
        """
        try:
            label = image_data.get('label', f"IMG_{image_data['paragraph_index']}")
            para_index = image_data['paragraph_index']
            
            # Crea bookmark univoco
            bookmark_name = self._sanitize_bookmark_name(f"DESC_{label}")
            
            # Inserisci link sotto l'immagine
            if para_index < len(self.document.paragraphs):
                img_paragraph = self.document.paragraphs[para_index]
                new_para = self._insert_paragraph_after(img_paragraph)
                new_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                self._add_hyperlink(new_para, bookmark_name, f"📖 Descrizione {label}")
            
            # Crea sezione descrizione se non esiste
            if self.descriptions_start_index is None:
                self.create_descriptions_section()
            
            # Aggiungi separatore
            sep_para = self.document.add_paragraph()
            sep_para.add_run("─" * 50).font.color.rgb = RGBColor(200, 200, 200)
            
            # Titolo descrizione con bookmark
            desc_title = self.document.add_paragraph()
            title_run = desc_title.add_run(f"📷 {label}")
            title_run.bold = True
            title_run.font.size = Pt(14)
            title_run.font.color.rgb = RGBColor(0, 102, 204)
            self._create_bookmark(desc_title, bookmark_name)
            
            # Testo descrizione
            desc_para = self.document.add_paragraph()
            desc_run = desc_para.add_run(description)
            desc_run.font.size = Pt(11)
            
            # Link di ritorno
            return_para = self.document.add_paragraph()
            return_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            return_run = return_para.add_run(return_text)
            return_run.font.size = Pt(9)
            return_run.font.color.rgb = RGBColor(128, 128, 128)
            return_run.italic = True
            
            self.document.add_paragraph()
            
            # Salva nel registro
            jump_info = {
                'type': 'image',
                'label': label,
                'source_paragraph': para_index,
                'bookmark': bookmark_name,
                'description_length': len(description)
            }
            
            self.jump_map[label] = jump_info
            return jump_info
            
        except Exception as e:
            print(f"Errore creazione jump: {e}")
            return None
    
    def create_descriptions_section(self, title="📚 DESCRIZIONI IMMAGINI E FIGURE"):
        """Crea la sezione delle descrizioni alla fine del documento"""
        self.document.add_page_break()
        
        title_para = self.document.add_paragraph()
        title_run = title_para.add_run(title)
        title_run.bold = True
        title_run.font.size = Pt(16)
        title_run.font.color.rgb = RGBColor(0, 51, 102)
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        subtitle = self.document.add_paragraph()
        subtitle_run = subtitle.add_run(
            "Clicca sui link 'Descrizione' nel testo per saltare qui.\n"
            "Clicca '↑ Torna al testo' per tornare alla posizione originale."
        )
        subtitle_run.font.size = Pt(10)
        subtitle_run.font.italic = True
        subtitle_run.font.color.rgb = RGBColor(100, 100, 100)
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        self.document.add_paragraph()
        self.descriptions_start_index = len(self.document.paragraphs)
        
        return self.descriptions_start_index
    
    def add_simple_description(self, label, description, image_path=None):
        """Aggiunge una descrizione semplice"""
        try:
            sep = self.document.add_paragraph()
            sep.add_run("─" * 50).font.color.rgb = RGBColor(200, 200, 200)
            
            title = self.document.add_paragraph()
            bookmark_name = self._sanitize_bookmark_name(f"DESC_{label}")
            
            title_run = title.add_run(f"📷 {label}")
            title_run.bold = True
            title_run.font.size = Pt(14)
            title_run.font.color.rgb = RGBColor(0, 102, 204)
            self._create_bookmark(title, bookmark_name)
            
            if image_path:
                try:
                    img_para = self.document.add_paragraph()
                    img_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    img_para.add_run().add_picture(image_path, width=Inches(3))
                except:
                    pass
            
            desc = self.document.add_paragraph()
            desc.add_run(description).font.size = Pt(11)
            
            ret = self.document.add_paragraph()
            ret.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            ret_run = ret.add_run("↑ Torna al testo")
            ret_run.font.size = Pt(9)
            ret_run.font.color.rgb = RGBColor(128, 128, 128)
            ret_run.italic = True
            
            self.document.add_paragraph()
            
            # Salva nel registro
            jump_info = {
                'type': 'simple',
                'label': label,
                'bookmark': bookmark_name,
                'description_length': len(description)
            }
            
            self.jump_map[label] = jump_info
            return jump_info
            
        except Exception as e:
            print(f"Errore add_simple_description: {e}")
            return None
    
    def _insert_paragraph_after(self, paragraph):
        """Inserisce un nuovo paragrafo dopo quello specificato"""
        new_p = OxmlElement('w:p')
        paragraph._element.addnext(new_p)
        from docx.text.paragraph import Paragraph
        return Paragraph(new_p, paragraph._parent)
    
    def _add_hyperlink(self, paragraph, bookmark_name, text):
        """Aggiunge un hyperlink a un bookmark"""
        p = paragraph._element
        hyperlink = OxmlElement('w:hyperlink')
        hyperlink.set(qn('w:anchor'), bookmark_name)
        
        r = OxmlElement('w:r')
        rPr = OxmlElement('w:rPr')
        
        color = OxmlElement('w:color')
        color.set(qn('w:val'), '0066CC')
        rPr.append(color)
        
        u = OxmlElement('w:u')
        u.set(qn('w:val'), 'single')
        rPr.append(u)
        
        b = OxmlElement('w:b')
        rPr.append(b)
        
        r.append(rPr)
        
        t = OxmlElement('w:t')
        t.text = text
        r.append(t)
        
        hyperlink.append(r)
        p.append(hyperlink)
    
    def _create_bookmark(self, paragraph, bookmark_name):
        """Crea un bookmark in un paragrafo"""
        bookmark_id = str(len(self.jump_map))
        p = paragraph._element
        
        bookmark_start = OxmlElement('w:bookmarkStart')
        bookmark_start.set(qn('w:id'), bookmark_id)
        bookmark_start.set(qn('w:name'), bookmark_name)
        
        bookmark_end = OxmlElement('w:bookmarkEnd')
        bookmark_end.set(qn('w:id'), bookmark_id)
        
        p.insert(0, bookmark_start)
        p.append(bookmark_end)
    
    def _sanitize_bookmark_name(self, name):
        """Sanitizza il nome del bookmark"""
        sanitized = re.sub(r'[^a-zA-Z0-9_]', '_', name)
        if sanitized and not sanitized[0].isalpha():
            sanitized = 'BM_' + sanitized
        return sanitized or 'bookmark'
    
    def get_summary(self):
        """Ottieni riepilogo dei jump creati"""
        return {
            'total_jumps': len(self.jump_map),
            'jumps': list(self.jump_map.values()),
            'descriptions_section_index': self.descriptions_start_index
        }


def extract_images_from_document(doc_path):
    """Estrae tutte le immagini da un documento Word"""
    from docx import Document
    
    doc = Document(doc_path)
    images = []
    
    for i, paragraph in enumerate(doc.paragraphs):
        for run in paragraph.runs:
            inline_shapes = run._element.findall('.//{*}blip')
            
            for shape in inline_shapes:
                rId = shape.get('{http://schemas.openxmlformats.org/officeDocument/2006/relationships}embed')
                
                if rId:
                    image_part = doc.part.related_parts[rId]
                    label = find_image_label(doc, i)
                    
                    image_info = {
                        'paragraph_index': i,
                        'label': label,
                        'rId': rId,
                        'content_type': image_part.content_type,
                        'image_part': image_part,
                        'context_before': get_context(doc, i, -2),
                        'context_after': get_context(doc, i, 2)
                    }
                    
                    images.append(image_info)
    
    return images


def find_image_label(doc, para_index, search_range=3):
    """Cerca la label di un'immagine nei paragrafi circostanti"""
    patterns = [
        r'Fig\.?\s*(\d+\.?\d*)',
        r'Figura\.?\s*(\d+\.?\d*)',
        r'Figure\.?\s*(\d+\.?\d*)',
        r'Immagine\.?\s*(\d+\.?\d*)',
        r'Grafico\.?\s*(\d+\.?\d*)',
        r'Tab\.?\s*(\d+\.?\d*)',
        r'Tabella\.?\s*(\d+\.?\d*)',
    ]
    
    start = max(0, para_index - search_range)
    end = min(len(doc.paragraphs), para_index + search_range + 1)
    
    for i in range(start, end):
        text = doc.paragraphs[i].text
        
        for pattern in patterns:
            match = re.search(pattern, text, re.IGNORECASE)
            if match:
                return match.group(0)
    
    return f"IMG_{para_index}"


def get_context(doc, para_index, offset):
    """Ottieni testo dei paragrafi a offset dalla posizione corrente"""
    target = para_index + offset
    
    if offset < 0:
        start = max(0, target)
        end = para_index
        paragraphs = doc.paragraphs[start:end]
    else:
        start = para_index + 1
        end = min(len(doc.paragraphs), para_index + offset + 1)
        paragraphs = doc.paragraphs[start:end]
    
    return ' '.join([p.text for p in paragraphs if p.text.strip()])
